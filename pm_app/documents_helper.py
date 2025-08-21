import json
import ast
from django.conf import settings
from dotenv import load_dotenv
import os
from openai import OpenAI, AuthenticationError, RateLimitError, APIConnectionError
from .models import Project, Outcome, Benefit, Deliverable, Task
from docx import Document
import re as _re2
from django.http import JsonResponse
from docx.shared import Pt
from datetime import date, timedelta as td, datetime as _dt



BUDGET_RE = _re2.compile(r'(?P<currency>[$£€])\s?(?P<amount>[\d,]+(?:\.\d+)?)', _re2.I)
YEAR_RE = _re2.compile(r'\b(20[2-9]\d|203\d)\b')

load_dotenv() 

# 3. Get the API key from the environment
api_key = os.getenv("OPENAI_API_KEY") 
MODEL = "gpt-4o"
client = None

# 4. Check if the key exists and initialize the client
if api_key:
    client = OpenAI(api_key=api_key)
else:
    print("Warning: OPENAI_API_KEY not found in .env file.")


def _strip_code_fences(s: str) -> str:
    s = (s or "").strip()
    if s.startswith("```"):
        lines = s.splitlines()
        if len(lines) >= 2 and lines[-1].strip().startswith("```"):
            return "\n".join(lines[1:-1]).strip()
    return s

def _coerce_obj(raw_text: str):
    txt = _strip_code_fences(raw_text)
    try:
        obj = json.loads(txt)
        if not isinstance(obj, dict): raise ValueError
        return obj
    except Exception:
        obj = ast.literal_eval(txt)
        if not isinstance(obj, dict): raise ValueError
        return obj

def chat_call(messages, temperature=0.3) -> str:
    if not client:
        raise RuntimeError("OpenAI client is not initialized. Check API Key.")
    try:
        resp = client.chat.completions.create(model=MODEL, messages=messages, temperature=temperature)
        return resp.choices[0].message.content
    except (AuthenticationError, RateLimitError, APIConnectionError) as e:
        raise RuntimeError(f"OpenAI call failed: {e}") from e

def generate_comm_plan(desc: str) -> dict:
    system = f"""
    You are a senior project communications consultant.
    Create a Communication Plan JSON for the project described below.   
    <desc>{desc}</desc>

    Return ONLY a JSON object with this exact structure:
    {{
     "Objective":"...",
     "Stakeholders":[
       {{"Name":"...","Role":"...","CommunicationMethod":"Status Email / Standup / SteerCo / Board Pack",
         "Frequency":"Weekly / Fortnightly / Monthly / Ad-hoc","Responsible":"...","Priority":"High/Medium/Low",
         "PreferredDeliveryMethod":"Email / MS Teams / Slack / Portal","CommunicationGoal":"..."}}
     ],
     "Channels":["Email","MS Teams","Standup","SteerCo"],
     "Notes":"Short notes if any"
    }}
    Rules: Generate 8–12 relevant stakeholder rows. Tailor Roles & Frequency to the project description in <desc>.
    """
    raw_response = chat_call([{"role": "system", "content": system}])
    return _coerce_obj(raw_response)

# --- Safety Net Functions ---

def _priority_norm(v: str) -> str:
    v = (v or "").strip().lower()
    if v.startswith("h"): return "High"
    if v.startswith("l"): return "Low"
    return "Medium"

def _comm_row_from_dict(d: dict) -> dict:
    return {
        "Name": d.get("Name") or d.get("Stakeholder") or "",
        "Role": d.get("Role") or "",
        "CommunicationMethod": d.get("CommunicationMethod") or "Status Email",
        "Frequency": d.get("Frequency") or "Weekly",
        "Responsible": d.get("Responsible") or "Project Manager",
        "Priority": _priority_norm(d.get("Priority")),
        "PreferredDeliveryMethod": d.get("PreferredDeliveryMethod") or "Email",
        "CommunicationGoal": d.get("CommunicationGoal") or d.get("Purpose") or "",
    }

def _default_comm_from_facts(project_name: str) -> dict:
    # A sensible default plan if the AI call fails
    return {
        "Objective": f"Keep stakeholders for the '{project_name}' project aligned on schedule, risks, and go-live readiness.",
        "Stakeholders": [
            {"Name": "Project Manager", "Role":"Delivery Lead", "CommunicationMethod":"Daily Standup", "Frequency":"Daily", "Responsible":"Self", "Priority":"High", "PreferredDeliveryMethod":"MS Teams", "CommunicationGoal":"Coordinate delivery & unblock issues"},
            {"Name": "Executive Sponsor", "Role":"Sponsor", "CommunicationMethod":"Steering Committee", "Frequency":"Fortnightly", "Responsible":"Project Manager", "Priority":"High", "PreferredDeliveryMethod":"Board Pack / Email", "CommunicationGoal":"Secure decisions, manage risks"},
            {"Name": "Product Team", "Role":"Product", "CommunicationMethod":"Backlog Review", "Frequency":"Weekly", "Responsible":"Product Manager", "Priority":"High", "PreferredDeliveryMethod":"Jira / Teams", "CommunicationGoal":"Align on scope and priorities"},
            {"Name": "Tech Lead", "Role":"Technology", "CommunicationMethod":"Tech Sync", "Frequency":"Weekly", "Responsible":"Tech Lead", "Priority":"Medium", "PreferredDeliveryMethod":"Teams", "CommunicationGoal":"Resolve architectural issues"},
        ],
        "Channels": ["Email", "MS Teams", "Standup", "Steering Committee"],
        "Notes": "This is a default plan. The AI-generated plan could not be created."
    }

def normalize_comm_obj(comm_obj, project_name: str) -> dict:
    try:
        if not isinstance(comm_obj, dict):
            return _default_comm_from_facts(project_name)

        stakeholders_list = []
        raw_stakeholders = comm_obj.get("Stakeholders")
        if isinstance(raw_stakeholders, list) and raw_stakeholders:
            for item in raw_stakeholders:
                if isinstance(item, dict):
                    stakeholders_list.append(_comm_row_from_dict(item))

        if not stakeholders_list:
            return _default_comm_from_facts(project_name)

        return {
            "Objective": (comm_obj.get("Objective") or "").strip() or f"Communicate status, risks and decisions for {project_name}.",
            "Stakeholders": stakeholders_list,
            "Channels": comm_obj.get("Channels") or ["Email", "MS Teams", "Standup"],
            "Notes": comm_obj.get("Notes") or "",
        }
    except Exception:
        return _default_comm_from_facts(project_name)



#Financial Plan Generation

def generate_financial_plan(desc: str) -> dict:
    """
    Generates a financial plan with a data structure that perfectly matches the target screenshot.
    """
    # FINAL PROMPT VERSION
    system = f"""
    You are a senior financial planner following the PRINCE2 methodology.
    Based on the project description below, create a detailed Financial Plan JSON.
    <desc>{desc}</desc>

    Generate a JSON object with the exact keys: "summary", "stages", "expenses", "cashflow", "tolerance", and "governance".

    Follow these detailed instructions for each key:
    1.  **summary**: Write a 2-3 sentence overview of the project's financial objectives. This should be a single string of text.
    2.  **stages**: Create a list of 4-6 project stages. For each stage, provide a "name", "duration", and estimated "cost".
    3.  **expenses**: Create a list of 5-7 key expense items. For each item, provide a "category" and a "cost". This structure should be a list of dictionaries, like the 'stages' section.
        - Example: `{{ "category": "Staff Training", "cost": "£10000" }}`
    4.  **cashflow**: Create a dictionary of key financial metrics: "initial_investment", "monthly_outflow", "expected_return_on_investment_roi", and "break_even_point".
    5.  **tolerance**: Create a dictionary defining deviation limits for "time_tolerance", "cost_tolerance", and "quality_tolerance".
    6.  **governance**: Write a 2-3 sentence paragraph describing the project's financial governance, including review frequency and change control. This should be a single string of text, NOT a dictionary.

    Use the currency "£". Return ONLY the JSON object.
    """
    raw_response = chat_call([{"role":"system", "content":system}], 0.3)
    return _coerce_obj(raw_response)

def _rows_from_any(data):
    """Normalize various JSON shapes to a list of lists for a table."""
    if data is None or data == "":
        return None

    # Handle stringified JSON
    if isinstance(data, str):
        txt = _strip_code_fences(data)
        try: data = json.loads(txt)
        except Exception:
            try: data = ast.literal_eval(txt)
            except Exception: return [["Text"], [txt]]

    # Handle a list of dictionaries (most common case)
    if isinstance(data, list) and data and isinstance(data[0], dict):
        headers = list(data[0].keys())
        rows = [headers]
        for item in data:
            rows.append([item.get(h, "") for h in headers])
        return rows
    
    # Handle a list of lists
    if isinstance(data, list) and data and isinstance(data[0], list):
        return data

    # Handle a simple dictionary of key-value pairs
    if isinstance(data, dict):
        return [["Field", "Value"]] + [[k, v] for k, v in data.items()]
        
    return None



#new updates 
def _parse_budget_year(txt: str):
    txt = (txt or "")
    m = BUDGET_RE.search(txt)
    budget = m.group(0) if m else ""
    y = YEAR_RE.search(txt)
    year = y.group(0) if y else ""
    return budget, year


def _project_facts(project_id: int) -> dict:
    try:
        p = Project.objects.get(pk=project_id)
    except Project.DoesNotExist:
        return {"Project Name": f"Project {project_id}",
                "Project Manager": "Project Manager", "Executive Sponsor": "Executive Sponsor",
                "Start Date": "", "End Date": "", "Total Budget": "", "Target Year": "",
                "Board Cadence": "Fortnightly", "Highlight Frequency": "Weekly",
                "Regulators": "", "Suppliers": [], "Objectives": [], "Deliverables": []}
    objectives = [o.description for o in p.outcomes.all()]
    deliverables = []
    for o in p.outcomes.all():
        for b in o.benefits.all():
            for d in b.deliverables.all():
                if d.description:
                    deliverables.append(d.description)
    b_from_vision, y_from_vision = _parse_budget_year(getattr(p, "vision", "") or "")
    return {
        "Project Name": getattr(p, "name", f"Project {project_id}"),
        "Project Manager": getattr(p, "project_manager", "Project Manager"),
        "Executive Sponsor": getattr(p, "sponsor", "Executive Sponsor"),
        "Start Date": getattr(p, "start_date", "") or "",
        "End Date": getattr(p, "end_date", "") or "",
        "Total Budget": getattr(p, "total_budget", "") or b_from_vision,
        "Target Year": getattr(p, "target_year", "") or y_from_vision,
        "Board Cadence": getattr(p, "board_cadence", "Fortnightly"),
        "Highlight Frequency": getattr(p, "highlight_frequency", "Weekly"),
        "Regulators": getattr(p, "regulators", ""),
        "Suppliers": getattr(p, "suppliers", []) or [],
        "Objectives": objectives,
        "Deliverables": deliverables,
    }


def build_project_desc(facts: dict) -> str:
    parts = []
    for k in ("Project Name", "Project Manager", "Executive Sponsor", "Total Budget",
              "Start Date", "End Date", "Board Cadence", "Highlight Frequency", "Regulators"):
        v = facts.get(k)
        if v:
            parts.append(f"{k}: {v}")
    if facts.get("Objectives"):
        parts.append("Objectives: " + ", ".join(facts["Objectives"]))
    if facts.get("Deliverables"):
        parts.append("Deliverables: " + ", ".join(facts["Deliverables"]))
    if facts.get("Suppliers"):
        parts.append("Suppliers: " + ", ".join(facts["Suppliers"]))
    return " | ".join(parts)


def _duration_to_days(s, default=7):
    if not s:
        return default
    try:
        num, unit = s.split()
        num = int(num)
        return num * 7 if 'week' in unit.lower() else num
    except Exception:
        return default



def _docx_add_table(doc: Document, rows, header: bool = True):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=1 if header else 0, cols=cols)
    if header:
        hdr = table.rows[0].cells
        for i, val in enumerate(rows[0]):
            run = hdr[i].paragraphs[0].add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.bold = True
        data = rows[1:]
    else:
        data = rows
    for r in data:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            run = cells[i].paragraphs[0].add_run("" if val is None else str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)


def _parse_money(s: str) -> float | None:
    if not s:
        return None
    m = BUDGET_RE.search(str(s))
    try:
        return float(m.group("amount").replace(",", "")) if m else float(str(s).replace(",", "").replace("£", "").strip())
    except Exception:
        return None


def _infer_dates_from_tasks(project: Project):
    tasks = list(Task.objects.filter(deliverableID__benefitID__outcomeID__projectID=project))
    starts = [t.start_date for t in tasks if t.start_date]
    ends = [t.end_date for t in tasks if t.end_date]
    if starts and ends:
        return min(starts), max(ends)
    s = date.today()
    return s, s + td(days=240)


def _normalize_stages_for_doc(stages_data, project: Project):
    """Return rows for the Stages table (with default objectives if missing)."""

    def _parse_date_any(x):
        if isinstance(x, (date, _dt)):
            return x.date() if isinstance(x, _dt) else x
        if not x:
            return date.today()
        s = str(x).strip()
        for fmt in ("%Y-%m-%d", "%d-%b-%Y", "%d-%B-%Y", "%d/%m/%Y", "%Y/%m/%d"):
            try:
                return _dt.strptime(s, fmt).date()
            except Exception:
                pass
        try:
            return _dt.fromisoformat(s.replace("Z", "")).date()
        except Exception:
            return date.today()

    # Helpful defaults we’ll apply when objectives are missing
    default_objectives_by_index = [
        "Approve business case, define success criteria, and secure funding.",
        "Baseline scope, schedule, budget; identify risks; finalize resources.",
        "Execute work packages, test increments, and manage changes/risks.",
        "Handover & training, benefits tracking setup, and project closeout."
    ]

    def _fallback_rows():
        names = ["Initiation", "Planning", "Execution", "Closure"]
        s, e = _infer_dates_from_tasks(project)
        total = max(1, (e - s).days)
        cuts = [s,
                s + td(days=int(total * .10)),
                s + td(days=int(total * .35)),
                s + td(days=int(total * .90)),
                e]
        rows = []
        for i, n in enumerate(names):
            rows.append([
                n,
                cuts[i].strftime("%d-%b-%Y"),
                cuts[i + 1].strftime("%d-%b-%Y"),
                default_objectives_by_index[i] if i < len(default_objectives_by_index) else ""
            ])
        return rows

    rows = []
    if isinstance(stages_data, list) and stages_data:
        for i, it in enumerate(stages_data):
            n   = (it.get("name") or it.get("Name") or f"Stage {i+1}").strip()
            sd  = _parse_date_any(it.get("start_date") or it.get("Start Date"))
            ed  = _parse_date_any(it.get("end_date")   or it.get("End Date"))
            objs = it.get("objectives") or it.get("Objectives") or it.get("objective") or ""
            if isinstance(objs, list):
                objs = ", ".join(o for o in objs if o)
            if not objs:
                # Apply a sensible default by position
                objs = default_objectives_by_index[i] if i < len(default_objectives_by_index) else "Define tasks, deliverables, and acceptance criteria."
            rows.append([n, sd.strftime("%d-%b-%Y"), ed.strftime("%d-%b-%Y"), objs])

    if not rows:
        rows = _fallback_rows()

    return [["name", "start_date", "end_date", "objectives"]] + rows

def _expenses_from_deliverables(project: Project):
    cats = ["Market Research", "Product Development", "Digital Marketing", "Branding and Design",
            "Website / Platform", "Staffing & Training", "Logistics & Distribution"]
    ds = list(Deliverable.objects.filter(benefitID__outcomeID__projectID=project))
    base_each = max(1, len(ds)) * 25000
    return [(c, f"£{base_each:,.0f}")] * len(cats)


def _monthly_cashflow(project: Project, total_cost_guess: float | None):
    s, e = _infer_dates_from_tasks(project)
    months, cur = [], date(s.year, s.month, 1)
    while cur <= e:
        months.append(cur)
        cur = date(cur.year + (1 if cur.month == 12 else 0), 1 if cur.month == 12 else cur.month + 1, 1)
    if total_cost_guess is None:
        total_cost_guess = sum(_parse_money(v) or 0 for _, v in _expenses_from_deliverables(project)) or 1_000_000
    monthly = total_cost_guess / max(1, len(months))
    return months, monthly, total_cost_guess




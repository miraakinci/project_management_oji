from .documents_helper import _project_facts, build_project_desc, _docx_add_table, _normalize_stages_for_doc, _expenses_from_deliverables, _parse_money, _monthly_cashflow, generate_comm_plan, generate_financial_plan, normalize_comm_obj, _rows_from_any
from .helper import find_similar_projects, find_similar_teams, serialize_project_flow, validate_and_serialize_sample_project
from .openapi_client import generate_flow_from_vision, update_flow_with_llm
from django.shortcuts import render, redirect, get_object_or_404

from datetime import date, timedelta as _timedelta, datetime as _dt
from .models import Project, Outcome, Benefit, Deliverable, Task

from django.http import JsonResponse, FileResponse
from datetime import date, timedelta
from django.http import HttpResponse
from django.db import transaction
import plotly.express as px
import plotly.offline as op
from .forms import InputForm
import pandas as pd
import traceback

from docx import Document
from docx.shared import Pt

import base64
import csv
import json

from io import BytesIO 


import matplotlib
matplotlib.use("Agg")   
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.patches import Rectangle





def index(request):
    """
    Handles the initial prompt submission page.
    """
    if request.method == 'POST':
        form = InputForm(request.POST)
        if form.is_valid():
            prompt = form.cleaned_data['prompt']

            #Create new project with the vision from the prompt
            project = Project.objects.create(name='initial project', vision=prompt)
            

            #find relevant past projects and organizational team/data
            similar_projects = find_similar_projects(prompt)
            teams_data = find_similar_teams(project.vision)

            # Validate and serialize the sample project into a clean JSON string
            sample_project_json = validate_and_serialize_sample_project(similar_projects)


            #generate the initial project flow
            project_flow_data = generate_flow_from_vision(prompt, sample_project_json, teams_data)

            #Update the project with the name from the flow data
            project_title = project_flow_data.get('title', 'Untitled Project')
            project.name = project_title
            project.save()

            # 3. Create related objects in the database
            for outcome_data in project_flow_data.get('outcomes', []):
                outcome = Outcome.objects.create(
                    projectID=project,
                    description=outcome_data.get('description')
                )
                #Create related benefits to each outcome 
                for benefit_data in outcome_data.get('benefits', []):
                    benefit = Benefit.objects.create(
                        outcomeID=outcome,
                        description=benefit_data.get('description')
                    )
                    #Create related deliverables to each benefit
                    for deliverable_data in benefit_data.get('deliverables', []):
                        deliverable = Deliverable.objects.create(
                            benefitID=benefit,
                            description=deliverable_data.get('description')
                        )
                        #Create related tasks to each deliverable
                        for task_data in deliverable_data.get('tasks', []):
                            start_date = task_data.get('start_date')
                            end_date = task_data.get('end_date')

                            Task.objects.create(
                                deliverableID=deliverable,
                                name=task_data.get('name'),
                                responsible_team=task_data.get('responsible_team', 'Unassigned'),
                                duration=task_data.get('duration', '1 day'),
                                start_date=date.fromisoformat(start_date) if start_date else None,
                                end_date=date.fromisoformat(end_date) if end_date else None
                            )
            # Redirect to the editable project flow page
            return redirect('project_flow', project_id=project.id)
        else:
            return render(request, 'pm_app/index.html', {'form': form, 'error': 'Invalid form submission'})
            
    return render(request, 'pm_app/index.html')




def get_project_flow(request, project_id):
    """
    Handles the page where the user can see and edit the project flow.
    """
    project = get_object_or_404(Project, id=project_id)
    outcomes = project.outcomes.all()
    benefits = Benefit.objects.filter(outcomeID__projectID=project)
    deliverables = Deliverable.objects.filter(benefitID__outcomeID__projectID=project)
    tasks = Task.objects.filter(deliverableID__benefitID__outcomeID__projectID=project) 

    return render(request, "pm_app/project_flow.html", {
        "project": project,
        "outcomes": outcomes,
        "benefits": benefits,
        "deliverables": deliverables,
        "tasks": tasks, 
    })



def update_flow_ajax(request, project_id):
    """
    Handles AJAX requests for project flow updates. It retrieves the entire current
    project flow from the database, combines it with the user's edit, and sends
    the full context to an LLM to generate the complete, updated flow.
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            project = get_object_or_404(Project.objects.prefetch_related(
                'outcomes__benefits__deliverables__tasks'
            ), id=project_id)
            
            with transaction.atomic():
                edited_field = data.get('edited_field')
                payload = data.get('payload')
                
                if not edited_field or not payload:
                    return JsonResponse({'status': 'error', 'message': 'Missing edited_field or payload.'}, status=400)


                # First, update the specific item in the database so the change is reflected.
                if edited_field == 'vision':
                    project.vision = payload.get('vision')
                    
                elif edited_field == 'outcomes':
                    item = get_object_or_404(Outcome, id=payload.get('id'))
                    item.description = payload.get('description')
                    item.save()
                elif edited_field == 'benefits':
                    item = get_object_or_404(Benefit, id=payload.get('id'))
                    item.description = payload.get('description')
                    item.save()
                elif edited_field == 'deliverables':
                    item = get_object_or_404(Deliverable, id=payload.get('id'))
                    item.description = payload.get('description')
                    item.save()
                
                project.save() # Save any changes to the project model itself (like vision)

                # Now, serialize the fully updated project to send to the LLM
                current_project_flow = serialize_project_flow(project)
                
                similar_projects = find_similar_projects(project.vision)
                similar_teams = find_similar_teams(project.vision)

                llm_payload = {
                    'edited_field': edited_field,
                    'user_edit': payload,
                    'current_flow': current_project_flow,
                    'similar_projects': similar_projects,
                    'similar_teams': similar_teams,
                }
                
                # Call the LLM to get the complete, re-aligned project flow
                updated_flow_data = update_flow_with_llm(llm_payload)
                
                if not updated_flow_data:
                    return JsonResponse({'status': 'error', 'message': 'LLM failed to return valid data.'}, status=500)

                # Clear old data and repopulate with the new, LLM-generated data
                project.outcomes.all().delete()
                project.name = updated_flow_data.get('title', project.name)
                # The vision is already up-to-date, but we save it again with the title
                project.save()

                for outcome_data in updated_flow_data.get('outcomes', []):
                    outcome = Outcome.objects.create(
                        projectID=project, description=outcome_data.get('description'))
                    for benefit_data in outcome_data.get('benefits', []):
                        benefit = Benefit.objects.create(
                            outcomeID=outcome, description=benefit_data.get('description'))
                        for deliverable_data in benefit_data.get('deliverables', []):
                            deliverable = Deliverable.objects.create(
                                benefitID=benefit, description=deliverable_data.get('description'))
                            for task_data in deliverable_data.get('tasks', []):
                                Task.objects.create(
                                    deliverableID=deliverable, name=task_data.get('name'),
                                    responsible_team=task_data.get('responsible_team'),
                                    duration=task_data.get('duration'))
                
                return JsonResponse({'status': 'success', 'message': 'Project flow updated successfully.'})

        except Exception as e:
            traceback.print_exc()
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

    return JsonResponse({'status': 'error', 'message': 'Invalid request method.'}, status=405)




def _duration_to_days(s, default=7):
    if not s: return default
    try:
        num, unit = s.split()
        num = int(num)
        return num * 7 if 'week' in unit.lower() else num
    except Exception:
        return default
    



def gantt_chart_data(request, project_id):

    project = get_object_or_404(Project, id=project_id)
    qs = (Task.objects
          .filter(deliverableID__benefitID__outcomeID__projectID=project)
          .order_by('id'))
    if not qs.exists():
        return JsonResponse({"png": None, "message": "No tasks found."})

    # Build rows; guarantee each task has a positive span
    rows, rolling = [], date.today()
    for t in qs:
        if t.start_date and t.end_date:
            start, end = t.start_date, t.end_date
        else:
            days = _duration_to_days(t.duration)
            start = (t.start_date or rolling)
            end = start + timedelta(days=max(1, days))
            rolling = end + timedelta(days=1)
        if end <= start:
            end = start + timedelta(days=1)
        rows.append({"task": t.name or "Untitled Task",
                     "team": t.responsible_team or "Unassigned",
                     "start": start, "end": end})

    # shared color palette (available to both try/except paths)
    palette = ["#4A90E2", "#50E3C2", "#F5A623", "#D0021B", "#7B61FF", "#417505",
               "#B8E986", "#F8E71C", "#BD10E0", "#7ED321", "#9013FE", "#F56A79"]

    # Try Matplotlib → PNG
    try:
        n = len(rows)
        fig, ax = plt.subplots(figsize=(11, max(2.5, 0.8 * n + 1)))

        rows_sorted = sorted(rows, key=lambda r: (r["start"], r["end"], r["task"]))
        y_labels = [f"Task {i + 1}" for i, _ in enumerate(rows_sorted)]

        # x-scale
        ax.set_xlabel("Date")
        ax.xaxis.set_major_locator(mdates.MonthLocator())
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %Y"))
        ax.grid(True, axis="x", linestyle=":", alpha=.35)

        # draw bars (visible fill + edge)
        for i, r in enumerate(rows_sorted):
            s = mdates.date2num(r["start"])
            e = mdates.date2num(r["end"])
            w = max(0.25, e - s)
            color = palette[i % len(palette)]
            rect = Rectangle(
                (s, i + 0.2), w, 0.6,
                facecolor=color, edgecolor="#333", linewidth=0.8, alpha=0.9
            )
            ax.add_patch(rect)

        # y axis
        ax.set_yticks([i + 0.5 for i in range(len(y_labels))])
        ax.set_yticklabels(y_labels)
        ax.set_ylim(0, len(y_labels) + 0.5)  # ensure bars are within view

        # x limits
        start_min = min(r["start"] for r in rows_sorted)
        end_max   = max(r["end"]   for r in rows_sorted)
        ax.set_xlim(mdates.date2num(start_min), mdates.date2num(end_max))

        ax.set_title("Project Gantt Schedule")
        fig.tight_layout()

        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=170, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        b64 = base64.b64encode(buf.read()).decode("ascii")
        resp = JsonResponse({"png": f"data:image/png;base64,{b64}"})
    except Exception:
        # SVG fallback — also colored
        start_min = min(r["start"] for r in rows); end_max = max(r["end"] for r in rows)
        total_days = max(1, (end_max - start_min).days)

        W, H = 1100, 90 + 28 * len(rows)
        L, R, T, B = 140, 20, 40, 20

        def x_for(d): return L + int((d - start_min).days / total_days * (W - L - R))

        svg = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}"><rect width="100%" height="100%" fill="#f8f9fb"/>']
        cur = date(start_min.year, start_min.month, 1)
        while cur <= end_max:
            x = x_for(cur)
            svg.append(f'<line x1="{x}" y1="{T}" x2="{x}" y2="{H-B}" stroke="#ddd" stroke-dasharray="3,3"/>')
            svg.append(f'<text x="{x+4}" y="{T-8}" font-size="11" fill="#666">{cur.strftime("%b %Y")}</text>')
            cur = date(cur.year + (1 if cur.month == 12 else 0), 1 if cur.month == 12 else cur.month + 1, 1)
        for i, r in enumerate(rows):
            y = T + 20 + i*28; x1 = x_for(r["start"]); x2 = x_for(r["end"])
            color = palette[i % len(palette)]
            svg.append(f'<rect x="{x1}" y="{y}" width="{max(2,x2-x1)}" height="14" fill="{color}" stroke="#333" stroke-width="1" rx="3" ry="3"/>')
            svg.append(f'<text x="10" y="{y+12}" font-size="12" fill="#333">Task {i+1}</text>')
        svg.append("</svg>")
        b64 = base64.b64encode("".join(svg).encode()).decode()
        resp = JsonResponse({"png": f"data:image/svg+xml;base64,{b64}"})

    # expose number→name map to the page
    task_map = {f"Task {i + 1}": r["task"] for i, r in enumerate(rows)}
    resp["X-Task-Map"] = json.dumps(task_map, ensure_ascii=False)
    return resp






def download_comm_plan_docx(request, project_id: int):
    """
    Generate Communication Plan (DOCX) with Stakeholders + Channels sections.
    """
    facts = _project_facts(project_id)
    desc = build_project_desc(facts)

    try:
        comm_raw = generate_comm_plan(desc)
    except Exception:
        comm_raw = {}

    comm = normalize_comm_obj(comm_raw or {}, facts.get("Project Name", "Project"))

    doc = Document()
    doc.add_heading(f"Communication Plan – {facts.get('Project Name', 'Project')}", level=1)

    # Summary / Objective
    obj = comm.get("Objective") or "Ensure alignment and timely decisions across stakeholders."
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(obj)

    # Stakeholders table
    stakeholders = comm.get("Stakeholders") or []
    if stakeholders:
        headers = list(stakeholders[0].keys())
        rows = [headers] + [[s.get(h, "") for h in headers] for s in stakeholders]
        doc.add_heading("Stakeholders", level=2)
        _docx_add_table(doc, rows, header=True)

    # Channels list
    channels = comm.get("Channels") or []
    if channels:
        doc.add_heading("Channels", level=2)
        for ch in channels:
            doc.add_paragraph(ch, style="List Bullet")

    # Return DOCX
    from tempfile import NamedTemporaryFile
    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.flush()
    tmp.seek(0)
    return FileResponse(
        open(tmp.name, "rb"),
        as_attachment=True,
        filename=f"project_{project_id}_communication_plan.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



def download_financial_plan_docx(request, project_id: int):
    from tempfile import NamedTemporaryFile  # safe to leave here even if also imported at top

    project = get_object_or_404(Project, pk=project_id)
    facts = _project_facts(project_id)
    desc = build_project_desc(facts)

    # Try AI; fall back safely
    try:
        ai_fin = generate_financial_plan(desc) or {}
    except Exception:
        ai_fin = {}

    # --- Summary ---
    summary_text = ""
    s = ai_fin.get("summary")
    if isinstance(s, dict):
        summary_text = s.get("Text") or s.get("text") or ""
    elif isinstance(s, str):
        summary_text = s
    if not summary_text:
        summary_text = (
            f"The financial plan for {facts.get('Project Name','this project')} covers stages, "
            f"costs, monthly phasing, tolerances and governance. Values below include sensible defaults "
            f"if AI data was unavailable."
        )

    # --- Stages ---
    stages = _normalize_stages_for_doc(ai_fin.get("stages"), project)

    # --- Expenses ---
    # Accept either a list of dicts from AI or fall back to deliverable-based defaults
    expenses_rows = []
    expenses_obj = ai_fin.get("expenses") or ai_fin.get("costs") or ai_fin.get("Costs")
    if isinstance(expenses_obj, list) and expenses_obj and isinstance(expenses_obj[0], dict):
        headers = list(expenses_obj[0].keys())
        expenses_rows = [headers] + [[row.get(h, "") for h in headers] for row in expenses_obj]
    if not expenses_rows:
        expenses_rows = [["category", "cost"]] + [[c, v] for c, v in _expenses_from_deliverables(project)]

    # --- Cashflow ---
    total_cost = 0.0
    for r in expenses_rows[1:]:
        # second column assumed to be money-like "£123,456"
        val = _parse_money(r[1]) if len(r) > 1 else None
        total_cost += (val or 0.0)
    months, per_month, total_guess = _monthly_cashflow(project, total_cost if total_cost else None)
    cashflow_rows = [["month", "planned_outflow"]] + [[m.strftime("%b %Y"), f"£{per_month:,.0f}"] for m in months]

    # --- Build the DOCX ---
    doc = Document()
    doc.add_heading(f"Financial Plan – {facts.get('Project Name','Project')}", level=1)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(summary_text)

    doc.add_heading("Stages", level=2)
    _docx_add_table(doc, stages, header=True)

    doc.add_heading("Expenses", level=2)
    _docx_add_table(doc, expenses_rows, header=True)

    doc.add_heading("Cashflow – Monthly Phasing", level=2)
    _docx_add_table(doc, cashflow_rows, header=True)

    doc.add_heading("Tolerance", level=2)
    _docx_add_table(doc, [["Field", "Value"], ["time_tolerance", "10%"], ["cost_tolerance", "15%"], ["quality_tolerance", "5%"]], header=True)

    doc.add_heading("Governance", level=2)
    gov_text = (
        f"Executive Sponsor: {facts['Executive Sponsor']}; PM: {facts['Project Manager']}. "
        f"Board cadence: {facts['Board Cadence']}; highlights: {facts['Highlight Frequency']}."
    )
    _docx_add_table(doc, [["Text", gov_text]], header=True)

    # --- Return file ---
    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.flush()
    tmp.seek(0)
    return FileResponse(
        open(tmp.name, "rb"),
        as_attachment=True,
        filename=f"project_{project_id}_financial_plan.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

